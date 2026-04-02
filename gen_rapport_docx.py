"""
gen_rapport_docx.py — Rapport exécutif as Word (.docx)
Tijan AI — same data as generer_rapport_executif in gen_mep.py
"""
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


VERT = RGBColor(0x43, 0xA9, 0x56)
GRIS = RGBColor(0x55, 0x55, 0x55)


def _styled_table(doc, headers, rows, col_widths=None):
    """Add a table with green header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '43A956', qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def _fmt(n):
    """Format number with thousands separator."""
    if isinstance(n, (int, float)):
        return f'{n:,.0f}'
    return str(n)


def generer_rapport_executif_docx(rs, rm, params: dict) -> bytes:
    """Generate Rapport exécutif as Word. Returns docx bytes."""
    doc = Document()
    d = rs.params
    boq_s = rs.boq
    boq_m = rm.boq
    e = rm.edge

    # Title
    title = doc.add_heading(d.nom, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = VERT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Rapport de synthèse exécutif — {d.ville} — R+{d.nb_niveaux-1}')
    run.font.size = Pt(12)
    run.font.color.rgb = GRIS

    doc.add_paragraph(
        'Ce document est destiné au maître d\'ouvrage. '
        'Il présente les points clés du projet et l\'estimation budgétaire globale.'
    )

    # Section 1 — Fiche Projet
    doc.add_heading('1. FICHE PROJET', level=1)
    fiche_data = [
        ('Projet', d.nom, 'Localisation', f'{d.ville}, {d.pays}'),
        ('Usage', d.usage.value.capitalize(), 'Hauteur', f'R+{d.nb_niveaux-1} ({d.nb_niveaux} niveaux)'),
        ('Surface bâtie', f'{_fmt(boq_s.surface_batie_m2)} m²', 'Surface habitable', f'{_fmt(boq_s.surface_habitable_m2)} m²'),
        ('Logements', str(rm.nb_logements), 'Occupants', str(rm.nb_personnes)),
        ('Béton', rs.classe_beton, 'Fondation', rs.fondation.type.value),
        ('Certification EDGE', '✓ Certifiable' if e.certifiable else '✗ Non certifiable', 'Conformité EC2', rs.analyse.conformite_ec2),
    ]
    _styled_table(doc,
        ['Paramètre', 'Valeur', 'Paramètre', 'Valeur'],
        fiche_data,
        col_widths=[4, 4, 4, 4],
    )

    # Section 2 — Budget
    doc.add_heading('2. ESTIMATION BUDGÉTAIRE GLOBALE', level=1)
    doc.add_paragraph('Estimation ±15% — Prix unitaires marché local 2026').italic = True

    total_bas = boq_s.total_bas_fcfa + boq_m.total_basic_fcfa
    total_haut = boq_s.total_haut_fcfa + boq_m.total_hend_fcfa
    pct_s = boq_s.total_bas_fcfa / total_bas * 100 if total_bas else 0
    pct_m = boq_m.total_basic_fcfa / total_bas * 100 if total_bas else 0

    budget_data = [
        ('Structure (gros œuvre)', _fmt(boq_s.total_bas_fcfa), _fmt(boq_s.total_haut_fcfa), f'{pct_s:.0f}%', 'Béton + acier + fondations'),
        ('MEP — Basic', _fmt(boq_m.total_basic_fcfa), _fmt(boq_m.total_hend_fcfa), f'{pct_m:.0f}%', 'Élec, plomb, CVC, asc, sécu'),
        ('TOTAL GROS ŒUVRE', _fmt(total_bas), _fmt(total_haut), '100%', 'Hors finitions, VRD'),
        ('Finitions (est. 35%)', _fmt(int(total_bas*0.35)), _fmt(int(total_haut*0.35)), '~35%', 'Carrelage, peinture, etc.'),
        ('COÛT TOTAL ESTIMÉ', _fmt(int(total_bas*1.35)), _fmt(int(total_haut*1.35)), '', f'{int(total_bas/boq_s.surface_batie_m2):,} FCFA/m²'),
    ]
    _styled_table(doc,
        ['Corps d\'état', 'Montant bas (FCFA)', 'Montant haut', '% Total', 'Note'],
        budget_data,
        col_widths=[4, 3.5, 3.5, 2, 4],
    )

    # Section 3 — EDGE
    doc.add_heading('3. PERFORMANCE ENVIRONNEMENTALE (EDGE IFC)', level=1)
    edge_data = [
        ('Économie énergie', f'{e.economie_energie_pct:.0f}%', '≥ 20%', '✓' if e.economie_energie_pct >= 20 else '✗'),
        ('Économie eau', f'{e.economie_eau_pct:.0f}%', '≥ 20%', '✓' if e.economie_eau_pct >= 20 else '✗'),
        ('Économie matériaux', f'{e.economie_materiaux_pct:.0f}%', '≥ 20%', '✓' if e.economie_materiaux_pct >= 20 else '✗'),
        ('VERDICT', e.niveau_certification, '3/3 piliers', '✓ CERTIFIABLE' if e.certifiable else '✗ NON CERTIFIABLE'),
    ]
    _styled_table(doc, ['Pilier', 'Score', 'Seuil', 'Statut'], edge_data, col_widths=[5, 3, 3, 5])

    # Footer
    doc.add_paragraph('\n\nDocument généré par Tijan AI — tijan.ai').italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

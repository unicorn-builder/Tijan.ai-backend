"""
gen_note_docx.py — Note de calcul structure as Word (.docx)
Tijan AI — same data as gen_note_structure.py, Word output via python-docx
"""
import io
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


VERT = RGBColor(0x43, 0xA9, 0x56)
NOIR = RGBColor(0x11, 0x11, 0x11)
GRIS = RGBColor(0x55, 0x55, 0x55)


def _add_table(doc, headers, rows, col_widths=None):
    """Add a styled table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '43A956', qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            p = cell.paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run(str(val))
            run.font.size = Pt(8)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def generer_note_structure_docx(rs, params: dict) -> bytes:
    """Generate Note de calcul structure as Word. Returns docx bytes."""
    doc = Document()
    d = rs.params
    boq = rs.boq
    ana = rs.analyse

    # Title
    title = doc.add_heading(d.nom, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = VERT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Note de calcul structure — {d.ville} — R+{d.nb_niveaux-1} — {d.usage.value.capitalize()}')
    run.font.size = Pt(12)
    run.font.color.rgb = GRIS

    doc.add_paragraph('Calculs indicatifs ±15% — À vérifier par un ingénieur structure habilité.').italic = True

    # Section 1 — Données du projet
    doc.add_heading('1. DONNÉES DU PROJET', level=1)
    fiche_data = [
        ('Projet', d.nom, 'Localisation', d.ville),
        ('Usage', d.usage.value.capitalize(), 'Niveaux', f'R+{d.nb_niveaux-1} ({d.nb_niveaux})'),
        ('Surface bâtie', f'{boq.surface_batie_m2:,.0f} m²', 'Surface habitable', f'{boq.surface_habitable_m2:,.0f} m²'),
        ('Portées', f'{d.portee_min_m}–{d.portee_max_m} m', 'Travées', f'{d.nb_travees_x}×{d.nb_travees_y}'),
        ('Béton', f'{rs.classe_beton} — fck={rs.fck_MPa:.0f} MPa', 'Acier', f'{rs.classe_acier} — fyk={rs.fyk_MPa:.0f} MPa'),
        ('Sol admissible', f'{rs.pression_sol_MPa} MPa', 'Zone sismique', f'Zone {rs.zone_sismique} — ag={rs.sismique.ag_g}g'),
    ]
    _add_table(doc,
        ['Paramètre', 'Valeur', 'Paramètre', 'Valeur'],
        fiche_data,
        col_widths=[4, 4, 4, 4],
    )

    doc.add_paragraph(f'\n{ana.justification_materiaux}').italic = True

    # Section 2 — Hypothèses
    doc.add_heading('2. HYPOTHÈSES ET NORMES DE CALCUL', level=1)
    hyp_data = [
        ('Béton armé', 'Eurocode 2 — NF EN 1992-1-1', f'γc=1.5 — fcd={rs.fck_MPa/1.5:.1f} MPa'),
        ('Séismique', 'Eurocode 8 — NF EN 1998-1', f'Zone {rs.zone_sismique} — ag={rs.sismique.ag_g}g'),
        ('Charges perm. G', 'EC1 — NF EN 1991-1-1', f'{rs.charge_G_kNm2} kN/m²'),
        ('Charges var. Q', 'EC1 — NF EN 1991-1-1', f'{rs.charge_Q_kNm2} kN/m²'),
        ('Combinaison ELU', '1.35G + 1.5Q', f'{1.35*rs.charge_G_kNm2+1.5*rs.charge_Q_kNm2:.1f} kN/m²'),
        ('Fondations', 'EC7 + DTU 13.2', f'qadm={rs.pression_sol_MPa} MPa — {rs.fondation.type.value}'),
    ]
    _add_table(doc, ['Domaine', 'Norme', 'Valeur'], hyp_data, col_widths=[4, 6, 5])

    # Section 3 — Poteaux
    doc.add_heading('3. DESCENTE DE CHARGES — POTEAUX (EC2/EC8)', level=1)
    pot_headers = ['Niveau', 'NEd (kN)', 'Section', 'Nb bar.', 'Ø (mm)', 'ρ (%)', 'NRd (kN)', 'NEd/NRd', 'Vérif.']
    pot_rows = []
    for pt in rs.poteaux:
        pot_rows.append([
            pt.niveau,
            f'{pt.NEd_kN:.0f}',
            f'{pt.section_mm}×{pt.section_mm}',
            str(pt.nb_barres),
            str(pt.diametre_mm),
            f'{pt.taux_armature_pct:.2f}',
            f'{pt.NRd_kN:.0f}',
            f'{pt.ratio_NEd_NRd:.2f}',
            '✓' if pt.verif_ok else '✗',
        ])
    _add_table(doc, pot_headers, pot_rows)

    # Section 4 — Poutres
    doc.add_heading('4. DIMENSIONNEMENT POUTRES (EC2)', level=1)
    for pout in [rs.poutre_principale, rs.poutre_secondaire]:
        if pout is None:
            continue
        doc.add_heading(f'Poutre {pout.type} — portée {pout.portee_m} m', level=2)
        pout_data = [[
            str(pout.b_mm), str(pout.h_mm),
            f'{pout.As_inf_cm2:.1f}', f'{pout.As_sup_cm2:.1f}',
            f'HA{pout.etrier_diam_mm}', f'{pout.etrier_esp_mm}',
            str(pout.portee_m),
        ]]
        _add_table(doc, ['b (mm)', 'h (mm)', 'As inf', 'As sup', 'Étriers', 'Esp.', 'Portée'], pout_data)

    # Section 5 — Dalle
    doc.add_heading('5. DALLE (EC2)', level=1)
    dalle = rs.dalle
    doc.add_paragraph(f'Épaisseur: {dalle.epaisseur_mm} mm — Portée: {dalle.portee_m} m')

    # Section 6 — Fondations
    doc.add_heading('6. FONDATIONS', level=1)
    fond = rs.fondation
    doc.add_paragraph(f'Type: {fond.type.value} — Profondeur: {fond.profondeur_m} m')
    doc.add_paragraph(f'Justification: {fond.justification}')

    # Footer
    doc.add_paragraph('\n\nDocument généré par Tijan AI — tijan.ai').italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
